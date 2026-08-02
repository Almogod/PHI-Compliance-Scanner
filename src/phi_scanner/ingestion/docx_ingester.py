"""DOCX ingester — yields ``(cell_text, SourceLocation)`` for paragraphs and table cells in Word documents.

Preserves exact location metadata:
- Paragraphs: location column="para", row=paragraph_index
- Table cells: location sheet_name="Table N", column=col_index, row=row_index
"""
from __future__ import annotations

from pathlib import Path
from typing import Iterator

import docx
from openpyxl.utils import get_column_letter

from .base import SourceLocation


class DocxIngester:
    """Reads Microsoft Word (.docx) files and yields text chunks with location provenance."""

    def ingest(self, path: Path) -> Iterator[tuple[str, SourceLocation]]:
        doc = docx.Document(path)

        # 1. Ingest paragraphs
        for idx, p in enumerate(doc.paragraphs, start=1):
            text = p.text.strip()
            if text:
                yield text, SourceLocation(
                    file_path=path,
                    sheet_name=None,
                    row=idx,
                    column="paragraph",
                )

        # 2. Ingest tables
        for t_idx, table in enumerate(doc.tables, start=1):
            sheet_name = f"Table {t_idx}"
            for r_idx, row in enumerate(table.rows, start=1):
                for c_idx, cell in enumerate(row.cells, start=1):
                    text = cell.text.strip()
                    if text:
                        col_letter = get_column_letter(c_idx)
                        yield text, SourceLocation(
                            file_path=path,
                            sheet_name=sheet_name,
                            row=r_idx,
                            column=col_letter,
                        )
