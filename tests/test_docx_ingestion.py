"""Tests for Microsoft Word (.docx) document ingestion."""
from pathlib import Path
import docx

from phi_scanner.engine import ScanEngine
from phi_scanner.recognizers.aadhaar import verhoeff_check_digit
from phi_scanner.recognizers.gstin import _gstin_check_digit


def make_aadhaar(prefix: str) -> str:
    return prefix + verhoeff_check_digit(prefix)


def make_gstin(prefix14: str) -> str:
    return prefix14 + _gstin_check_digit(prefix14)


def test_docx_paragraph_and_table_scanning(tmp_path: Path) -> None:
    docx_file = tmp_path / "contract.docx"

    doc = docx.Document()
    # Paragraph with Aadhaar & PAN
    valid_aadhaar = make_aadhaar("23456789012")
    doc.add_paragraph(f"Vendor Agreement for Aadhaar: {valid_aadhaar} and PAN: ABCPD1234E.")

    # Table with GSTIN & Mobile
    table = doc.add_table(rows=2, cols=2)
    valid_gstin = make_gstin("27ABCPD1234E1Z")
    table.cell(0, 0).text = "Vendor Name"
    table.cell(0, 1).text = "GSTIN"
    table.cell(1, 0).text = "Acme Corp"
    table.cell(1, 1).text = valid_gstin

    doc.save(docx_file)

    findings = list(ScanEngine().scan_file(docx_file))
    entity_types = {f.entity_type for f in findings}

    assert "AADHAAR" in entity_types
    assert "PAN" in entity_types
    assert "GSTIN" in entity_types

    # Provenance verification
    aadhaar_f = next(f for f in findings if f.entity_type == "AADHAAR")
    assert aadhaar_f.location.column == "paragraph"
    assert aadhaar_f.location.row == 1

    gstin_f = next(f for f in findings if f.entity_type == "GSTIN")
    assert gstin_f.location.sheet_name == "Table 1"
    assert gstin_f.location.row == 2
